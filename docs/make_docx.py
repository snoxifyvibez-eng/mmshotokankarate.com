# -*- coding: utf-8 -*-
"""Generate editable Word versions of the MMS Shotokan client documents."""
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERMILION = RGBColor(0xC0, 0x27, 0x1F)
INK = RGBColor(0x22, 0x24, 0x28)
MUTED = RGBColor(0x5B, 0x61, 0x6A)
SERIF = "Georgia"
SANS = "Calibri"


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def base_style(doc):
    st = doc.styles['Normal']
    st.font.name = SANS
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    # page margins
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)


def para(doc, text="", size=10.5, color=INK, bold=False, italic=False,
         font=SANS, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.bold = bold
        r.italic = italic
    return p


def eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text.upper())
    r.font.name = SANS
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = VERMILION
    # letter spacing
    rpr = r._element.get_or_add_rPr()
    spc = OxmlElement('w:spacing'); spc.set(qn('w:val'), '30'); rpr.append(spc)
    return p


def heading(doc, text, size=15, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = SERIF
    r.font.size = Pt(size)
    r.bold = True
    r.font.color.rgb = INK
    return p


def section_num(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(num + "   ")
    r1.font.name = SERIF; r1.font.size = Pt(11); r1.bold = True; r1.font.color.rgb = VERMILION
    r2 = p.add_run(title)
    r2.font.name = SERIF; r2.font.size = Pt(14); r2.bold = True; r2.font.color.rgb = INK
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.name = SANS; r.font.size = Pt(10.5)
    r = p.add_run(text); r.font.name = SANS; r.font.size = Pt(10.5)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.name = SANS; r.font.size = Pt(10.5)
    return p


def hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'C0271F')
    pbdr.append(bottom); pPr.append(pbdr)


def make_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], '1B1E23')
        pg = hdr[i].paragraphs[0]; pg.paragraph_format.space_after = Pt(2)
        r = pg.add_run(h.upper()); r.font.name = SANS; r.font.size = Pt(8); r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pg = cells[i].paragraphs[0]; pg.paragraph_format.space_after = Pt(2)
            # bold marker: prefix with '**'
            bold = False
            if isinstance(val, str) and val.startswith('**'):
                bold = True; val = val[2:]
            r = pg.add_run(val); r.font.name = SANS; r.font.size = Pt(9.5); r.bold = bold
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)
    return t


def letterhead(doc, meta_lines):
    t = doc.add_table(rows=1, cols=2)
    left, right = t.rows[0].cells
    # left: seal + name
    p = left.paragraphs[0]
    r = p.add_run("押  "); r.font.name = SERIF; r.font.size = Pt(20); r.bold = True; r.font.color.rgb = VERMILION
    r = p.add_run("[Your Name / Studio]"); r.font.name = SERIF; r.font.size = Pt(13); r.bold = True
    p2 = left.add_paragraph(); r = p2.add_run("Web Design & Development · Toronto, ON")
    r.font.name = SANS; r.font.size = Pt(9); r.font.color.rgb = MUTED
    # right meta
    for i, line in enumerate(meta_lines):
        pr = right.paragraphs[0] if i == 0 else right.add_paragraph()
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = pr.add_run(line); r.font.name = SANS; r.font.size = Pt(9); r.font.color.rgb = MUTED
    # remove borders
    for cell in (left, right):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'nil'); borders.append(e)
        tcPr.append(borders)
    hrule(doc)


def footer(doc):
    hrule(doc)
    para(doc, "[Your Name / Studio]  ·  [email]  ·  [phone]", size=8.5, color=MUTED, space_before=4)
    para(doc, "押忍 — MMS Shotokan Karate", size=8.5, color=VERMILION, font=SERIF)


# ============================================================ PROPOSAL
def build_proposal(path):
    doc = Document(); base_style(doc)
    letterhead(doc, ["Proposal #2026-014", "Date: July 10, 2026", "Valid until: August 9, 2026"])

    para(doc, "Website Design, Registration & Online Payments", size=22, font=SERIF, bold=True, space_before=8, space_after=4)
    p = doc.add_paragraph()
    r = p.add_run("Prepared for "); r.font.name = SANS; r.font.size = Pt(10.5)
    r = p.add_run("MMS Shotokan Karate"); r.bold = True; r.font.name = SANS; r.font.size = Pt(10.5)
    r = p.add_run(" — Oshawa, Ontario   ·   Attn: "); r.font.name = SANS; r.font.size = Pt(10.5)
    r = p.add_run("[Dojo Owner / Sensei]"); r.bold = True; r.font.name = SANS; r.font.size = Pt(10.5)

    para(doc, "A custom, mobile-friendly website for the dojo that does more than look good: prospective "
              "students can register online, families can pay the registration fee and monthly membership "
              "securely, and every enquiry lands straight in your inbox — so you spend less time on admin and "
              "more on the mat.", space_before=8)

    section_num(doc, "01", "What this project delivers")
    for label, txt in [("Goal — ", "Turn website visitors into registered, paying students without manual paperwork."),
                       ("For — ", "Parents and adults in Oshawa/Durham searching for traditional Shotokan karate."),
                       ("Approach — ", "A bespoke design — clean, disciplined, and unmistakably a karate dojo."),
                       ("Result — ", "Registrations by email + automated Stripe payments, on your own hosting.")]:
        bullet(doc, txt, bold_prefix=label)

    section_num(doc, "02", "Scope & deliverables")
    groups = [
        ("Custom website & design", ["Bespoke visual identity & responsive layout (phones, tablets, desktop)",
                                       "Five pages: Home, Contact, Pricing, plus payment success & cancel pages",
                                       "Full content: schedule, location, curriculum, values, who can join",
                                       "Smooth animations, accessibility-minded markup, fast loading"]),
        ("Online registration & contact", ["Registration & contact forms that email submissions to your Gmail",
                                            "Spam protection built in (honeypot + trusted form service)",
                                            "No monthly software fee — uses a free, reliable form service"]),
        ("Secure online payments (Stripe)", ["One-time registration fee + recurring monthly membership",
                                              "Secure Stripe Checkout — cards, Apple Pay & Google Pay",
                                              "Automatic email notification to the dojo on every payment",
                                              "Members can cancel/manage their own membership",
                                              "Card data never touches your site"]),
        ("Launch & handover", ["Deployment to your Hostinger hosting with HTTPS",
                               "Written setup guide for keys, pricing & future edits",
                               "End-to-end testing of forms & payments before go-live"]),
    ]
    for title, items in groups:
        heading(doc, title, size=11.5, space_before=8)
        for it in items:
            bullet(doc, it)

    section_num(doc, "03", "Investment")
    make_table(doc, ["Work", "Est. effort", "Amount (CAD)"], [
        ["Discovery & custom design", "~14 hrs", "$1,400"],
        ["Website build (5 pages)", "~18 hrs", "$1,800"],
        ["Registration & contact email", "~4 hrs", "$400"],
        ["Stripe payment platform", "~14 hrs", "$1,700"],
        ["Deployment, testing & QA", "~6 hrs", "$600"],
        ["Project management & revisions", "~8 hrs", "$600"],
        ["**Fixed project total", "**~64 hrs", "**$6,500"],
    ], widths=[3.6, 1.2, 1.4])
    para(doc, "Quoted as a fixed price, not hourly — hours shown for transparency. Stripe's processing fee "
              "(≈ 2.9% + 30¢ per transaction) is charged by Stripe to the dojo and is not part of this quote.",
         size=8.5, color=MUTED, space_before=6)

    section_num(doc, "04", "Ongoing care & optional add-ons")
    p = doc.add_paragraph()
    r = p.add_run("Recommended care plan — "); r.bold = True; r.font.name = SANS; r.font.size = Pt(10.5)
    r = p.add_run("$85 / month"); r.bold = True; r.font.name = SERIF; r.font.size = Pt(12); r.font.color.rgb = VERMILION
    para(doc, "Hosting checks, updates, security, Stripe/plugin maintenance, and small content edits "
              "(up to 1 hr/mo). Cancel anytime.", size=9.5, color=MUTED, space_before=0)
    make_table(doc, ["Optional add-on", "Price"], [
        ["Family / multi-student pricing tier", "+$450"],
        ["Logo & brand mark design", "+$600"],
        ["Professional copywriting (all pages)", "+$400"],
        ["Photography of classes & dojo", "Referral (local partner)"],
    ], widths=[4.4, 1.8])

    section_num(doc, "05", "Timeline")
    make_table(doc, ["Phase", "Work", "When"], [
        ["1 · Kickoff & content", "Gather logo, photos, fees, Stripe/Hostinger access", "Week 1"],
        ["2 · Design & build", "Pages built, first preview link for feedback", "Weeks 1–2"],
        ["3 · Payments & forms", "Stripe + email wired and tested in test mode", "Week 3"],
        ["4 · Revisions & launch", "Final edits, go-live on Hostinger, handover", "Week 4"],
    ], widths=[1.9, 3.3, 1.0])
    para(doc, "Typical delivery: 3–4 weeks from kickoff, assuming content and account access are provided promptly.",
         size=8.5, color=MUTED, space_before=6)

    section_num(doc, "06", "Payment terms & assumptions")
    numbered(doc, "50% deposit ($3,250) to begin; 50% balance ($3,250) due on launch.")
    numbered(doc, "Care plan billed monthly, starting the month after launch. Cancel with 30 days' notice.")
    numbered(doc, "Includes two rounds of revisions; further changes billed at [$90]/hr or via the care plan.")
    numbered(doc, "Payment by e-transfer to [email] or invoice via [Stripe/PayPal].")
    for a in ["Dojo provides content (text, logo, photos) and Stripe + Hostinger account access.",
              "Stripe and Hostinger accounts are owned by the dojo; site ownership transfers on final payment.",
              "Third-party fees (hosting renewal, domain, Stripe processing) are the dojo's responsibility."]:
        bullet(doc, a)

    section_num(doc, "07", "Acceptance")
    para(doc, "To proceed, sign below and return with the 50% deposit. We'll schedule kickoff within a few days.",
         size=9.5, color=MUTED)
    para(doc, "")
    para(doc, "______________________________            ______________________________", space_before=18)
    para(doc, "Signature — MMS Shotokan Karate                 Signature — [Your Name / Studio]", size=9, color=MUTED)
    para(doc, "Name & date                                                    Name & date", size=9, color=MUTED, space_before=14)

    footer(doc)
    doc.save(path)
    print("wrote", path)


# ============================================================ SETUP GUIDE
def build_setup(path):
    doc = Document(); base_style(doc)
    letterhead(doc, ["Setup Guide", "Project: MMS Shotokan website", "Date: July 10, 2026"])

    para(doc, "What we need to switch the website fully on", size=20, font=SERIF, bold=True, space_before=8, space_after=4)
    para(doc, "The website is built and ready. To finish the job, we need to connect a few outside services — "
              "the ones that send your enquiry emails and take payments. Most are free; Stripe (for payments) "
              "needs your business and bank details, which only you can provide. This guide explains each one "
              "in plain language and tells you exactly what to send us.", space_before=6)

    section_num(doc, "01", "The short version")
    make_table(doc, ["Service", "What it does", "Cost", "Who sets it up"], [
        ["**Hostinger hosting + domain", "Runs the website at your web address", "Already have", "You (access shared)"],
        ["**Web3Forms", "Emails you every contact & registration form", "Free", "You (2 min) → send us the key"],
        ["**Stripe", "Takes registration & membership payments", "Per-payment fee", "You create + verify → we connect"],
        ["**Gmail (dojo inbox)", "Where enquiries & payment alerts arrive", "Already have", "You — nothing to do"],
    ], widths=[1.7, 2.5, 1.1, 1.6])
    para(doc, "Stripe's fee is roughly 2.9% + C$0.30 per successful card payment, charged by Stripe (not us). "
              "There is no monthly fee.", size=8.5, color=MUTED, space_before=6)

    section_num(doc, "02", "Step by step")

    heading(doc, "A · Web3Forms — enquiry & registration emails  (Free)", size=11.5, space_before=8)
    para(doc, "Whenever someone fills the contact or registration form, this delivers it straight to your Gmail. "
              "No software to install.", size=9.5, color=MUTED)
    numbered(doc, "Go to web3forms.com.")
    numbered(doc, "Enter your dojo email: mms.shotokan.karate@gmail.com.")
    numbered(doc, "They email you an Access Key (a short code). Copy it.")
    p = doc.add_paragraph(); r = p.add_run("Send us: "); r.bold = True; r.font.color.rgb = VERMILION; r.font.name = SANS; r.font.size = Pt(10)
    r = p.add_run("the Access Key from your email. That's all — we plug it in."); r.font.name = SANS; r.font.size = Pt(10)

    heading(doc, "B · Stripe — online payments  (Per-payment fee)", size=11.5, space_before=8)
    para(doc, "Stripe securely handles the one-time registration fee and the monthly membership. Card details go "
              "straight to Stripe — they never touch your website. Money lands in your bank account.", size=9.5, color=MUTED)
    numbered(doc, "Create a free account at stripe.com using the dojo email.")
    numbered(doc, "Complete business verification — Stripe asks for dojo/owner details and ID (required by law).")
    numbered(doc, "Add the bank account where you want payouts deposited.")
    numbered(doc, "Decide your prices, e.g. Registration $[50] one-time and Membership $[80]/month, and tell us.")
    numbered(doc, "Invite us to your Stripe account (Settings → Team), or share the API keys with us securely.")
    p = doc.add_paragraph(); r = p.add_run("Send us (securely — never by plain email): "); r.bold = True; r.font.color.rgb = VERMILION; r.font.name = SANS; r.font.size = Pt(10)
    r = p.add_run("either a team invite, or your Secret key (sk_live_…) and the two Price IDs (price_…)."); r.font.name = SANS; r.font.size = Pt(10)

    heading(doc, "C · Hostinger — hosting & web address  (Already set up)", size=11.5, space_before=8)
    para(doc, "Your site lives here. We already have FTP access to publish updates. Two quick confirmations:", size=9.5, color=MUTED)
    numbered(doc, "Make sure the free SSL certificate is on (the padlock / https) — hPanel → SSL.")
    numbered(doc, "Confirm the site uses PHP 8 — hPanel → PHP Configuration.")

    heading(doc, "D · Dojo inbox  (Already have)", size=11.5, space_before=8)
    para(doc, "All enquiries, registrations, and payment alerts arrive at mms.shotokan.karate@gmail.com. "
              "Keep an eye on it (and the spam folder for the first few) after launch.", size=9.5, color=MUTED)

    section_num(doc, "03", "Your checklist — what to send us")
    for item in ["Web3Forms Access Key (from your email)",
                 "Stripe access — team invite, or Secret key + the two Price IDs",
                 "Your prices — registration (one-time) and membership (monthly) amounts",
                 "Confirm the domain and that SSL is on",
                 "Any content still to add — logo, photos, exact class details"]:
        bullet(doc, "☐  " + item)
    p = doc.add_paragraph()
    r = p.add_run("A note on security: "); r.bold = True; r.font.name = SANS; r.font.size = Pt(10)
    r = p.add_run("please don't send the Stripe Secret key or your FTP password in a normal email or text. "
                  "We'll share a safe method (a one-time secret link or a team invite). Treat those like the keys to your bank.")
    r.font.name = SANS; r.font.size = Pt(10)

    section_num(doc, "04", "Then what?")
    para(doc, "Once we have the items above, we connect everything, run a full test payment in Stripe's safe "
              "“test mode”, switch it to live, and publish. Typical turnaround is 2–3 business days after we "
              "receive your details. We'll confirm each step with you before going live.", size=9.5, color=MUTED)

    footer(doc)
    doc.save(path)
    print("wrote", path)


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "."
    build_proposal(out + "/proposal.docx")
    build_setup(out + "/setup-guide.docx")
